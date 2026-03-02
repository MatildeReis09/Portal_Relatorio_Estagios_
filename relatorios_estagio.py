import streamlit as st
import pandas as pd
from docx import Document 
from io import BytesIO
from datetime import datetime

# CONFIGURAÇÃO DA PÁGINA
st.set_page_config(page_title="Gerador de Relatórios", layout="wide")

# Link gerado pelas respostas do froms (google sheets) 
#URL_CSV_Respostas = "https://docs.google.com/spreadsheets/d/e/2PACX-1vSy4HVfG-ckMI2a0LHV17mq_HWnv5qwOj4R7yhcp7w_srg604hAWBSn6SrebKKI-cqrHKCcEwDo8jjA/pub?gid=1536866325&single=true&output=csv"
# link respetivo de ligação aos pins
#URL_Utilizadores = "https://docs.google.com/spreadsheets/d/e/2PACX-1vSy4HVfG-ckMI2a0LHV17mq_HWnv5qwOj4R7yhcp7w_srg604hAWBSn6SrebKKI-cqrHKCcEwDo8jjA/pub?gid=1984584767&single=true&output=csv"

#  nova forma de exportar dados excel diretamente , não ha delay 
ID_FOLHA = "1ELN40a8g8LOZkzBw20s2Yz3njWE5VunLE_gggqBxN2Y"
URL_CSV_Respostas = f"https://docs.google.com/spreadsheets/d/{ID_FOLHA}/export?format=csv&gid=1536866325"
URL_Utilizadores = f"https://docs.google.com/spreadsheets/d/{ID_FOLHA}/export?format=csv&gid=1984584767"

st.title("Portal de Relatórios de Estágio")
st.markdown("---")

# botão para forçar atualizar dados 
if st.sidebar.button("Atualizar Dados"):
    st.cache_data.clear()
    st.rerun()

def carregar_dados(url):
    try:
        # refresh=timestamp engana o cache do Google
        url_final = f"{url}&refresh={datetime.now().timestamp()}"
        df = pd.read_csv(url_final)
        df.columns = df.columns.str.strip()
        return df
     
    except Exception as e:  #aparecer o tipo de erro
        st.error(f"Erro ao carregar dados: {e}")
        return None
        

#criar e formata o documento para donwload
def create_word (dados, nome,titulo): 
    doc= Document()
    doc.add_heading(titulo, 0)
    doc.add_paragraph(f"Estagiário: {nome}")
    doc.add_paragraph(f"Data de Emissão:{datetime.now().strftime('%d/%m/%Y')}")

    # Ordenar por data de referência (trabalho)
    dados = dados.sort_values(by='Semana de Referência')
    
    for _, linha in dados.iterrows():
    # formata a datas para PT-PT no Word
        data_referencia = pd.to_datetime(linha['Semana de Referência']).strftime('%d/%m/%Y')
        doc.add_heading(f"Semana de {data_referencia}", level=1)

        p = doc.add_paragraph()
        p.add_run("Tarefas: ").bold = True
        p.add_run(str(linha['Tarefas Realizadas']))

        p = doc.add_paragraph()
        p.add_run("Aprendizagens: ").bold = True
        p.add_run(str(linha['Aprendizagens']))

        p = doc.add_paragraph()
        p.add_run("Dificuldades: ").bold = True
        p.add_run(str(linha['Dificuldades e Soluções']))

        doc.add_paragraph()
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer 

# Carregar dados
df_resp = carregar_dados(URL_CSV_Respostas)
df_auth = carregar_dados(URL_Utilizadores)


##interface de streamlit

if df_resp is not None and df_auth is not None: #as duas verdadeiras
    #sibebar parao login 
    st.sidebar.header("Autenticação de Utilizador")

    lista_nomes = sorted(df_auth['Utilizadores'].unique().tolist())
    nome_utilizador = st.sidebar.selectbox("Estagiário", lista_nomes)

     #criar o campo para o pin no sidebar e procurar o pin de cada utilizador
    pin = st.sidebar.text_input("Introduz o teu pin", type="password")
    
    # Validação do PIN
    user_row = df_auth[df_auth['Utilizadores'] == nome_utilizador]
    pin_correct = str(user_row['PIN'].values[0]).strip() if not user_row.empty else ""

    acesso_autorizado = False 
    if pin == pin_correct: 
        st.sidebar.success(f"Bem-vindo/a, {nome_utilizador}")
        acesso_autorizado = True

        ##st.write(f"A tentar encontrar: '{nome_utilizador}'")
        ##st.write("Nomes encontrados na folha:", df_resp['Nome'].unique().tolist())
    
    elif pin != "":
        st.sidebar.error("Pin Incorreto. Acesso bloqueado")
        
    else:
         st.sidebar.info("Insira o Pin para ter acesso aos seus dados de Relatórios")
        

if acesso_autorizado:

    #st.write("DEBUG: Nome que escolheste no login:", nome_utilizador)
    #st.write("DEBUG: Nomes que existem na folha de respostas:", df_resp['Nome'].unique().tolist())
   
    # Converter a data (dia/mês/ano do CSV costuma  para YYYY-MM-DD)
    df_resp['Carimbo de data/hora'] = pd.to_datetime(df_resp['Carimbo de data/hora'], errors='coerce')
    df_resp['Semana de Referência'] = pd.to_datetime(df_resp['Semana de Referência'], errors='coerce')

    tab_semanal, tab_mensal = st.tabs(["Atividades da semana", "Resumo de Atividades Mensal"])

    with tab_semanal: 

        # Filtrar pelo estagiario e datas 
        dados_estagiario = df_resp[df_resp['Nome'] == nome_utilizador.strip()]
        semanas_disponiveis = dados_estagiario.dropna(subset=['Semana de Referência'])
        
    
        # criar lista para selectbox de datas 
        lista_semanas = sorted(semanas_disponiveis['Semana de Referência'].dt.strftime('%d/%m/%Y').unique(), reverse=True)

        if lista_semanas: 
            semana_sel_str = st.selectbox("Escolha uma semana:", lista_semanas)
            
            # filtro de dados da semana escolhida
            dados_semanais = dados_estagiario[dados_estagiario['Semana de Referência'].dt.strftime('%d/%m/%Y') == semana_sel_str]

            if not dados_semanais.empty: 
                buffer_sem = create_word(dados_semanais, nome_utilizador, "Relatório Semanal")
                st.download_button( 
                    label="Descarregar Esta Semana (.docx)", 
                    data=buffer_sem, 
                    file_name=f"Semana_{nome_utilizador}_{semana_sel_str.replace('/','-')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else: 
                st.warning("Sem dados referentes a esta semana")

        else: 
            st.info("Ainda não existem registos para este utilizador")
                    ## vazia , sem tarefas



    with tab_mensal:
         #escolher o mes (já completo)
        mes_sel = st.slider("Selecione o Mês do Relatório", 1, 12, datetime.now().month)
        
        # Muito Importante ->  Filtramos pelo MÊS da Semana de Referência (data do trabalho)
        dados_mes = dados_estagiario[dados_estagiario['Semana de Referência'].dt.month == mes_sel]

        if not dados_mes.empty:
            #st.success(f"Encontrados {len(dados_mes)} registos para o mês {mes_sel}.")
            #st.dataframe(dados_mes[['Semana de Referência', 'Tarefas Realizadas']]) (tabela de visualização)
            
            buffer_mes = create_word(dados_mes, nome_utilizador, f"Relatório Mensal - Mês {mes_sel}")
            st.download_button(
                label="Descarregar Relatório Mensal (.docx)",
                data=buffer_mes, 
                file_name=f"Mensal_{nome_utilizador}_Mes_{mes_sel}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else: 
            st.warning("Ainda não existem dados para este mês.")









